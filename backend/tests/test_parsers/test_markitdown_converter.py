"""Tests for MarkitdownConverter and DirectFileImageHandler."""

import pytest
from pathlib import Path
from unittest.mock import MagicMock, patch
from backend.parsers.markitdown_converter import (
    MarkitdownConverter,
    MarkitdownConversionError,
    ConversionResult,
    ImageInfo,
    DirectFileImageHandler,
)


class TestMarkitdownConverter:

    def test_converter_initialization(self):
        converter = MarkitdownConverter()
        assert converter.timeout == 300

    def test_unsupported_file_type(self, tmp_path):
        unsupported_file = tmp_path / "document.txt"
        unsupported_file.write_text("dummy content")
        converter = MarkitdownConverter()
        with pytest.raises(ValueError) as exc_info:
            converter.convert(unsupported_file)
        assert "暂不支持" in str(exc_info.value)

    def test_nonexistent_file(self):
        converter = MarkitdownConverter()
        with pytest.raises(FileNotFoundError):
            converter.convert(Path("/nonexistent/document.docx"))

    @pytest.mark.integration
    def test_convert_sample_docx(self, tmp_path):
        pytest.skip("Requires sample DOCX file")

    def test_real_docx_materializes_images_and_never_emits_data_uri(self, tmp_path):
        from docx import Document
        from PIL import Image

        image_path = tmp_path / "fixture.png"
        Image.new("RGB", (640, 360), "white").save(image_path)
        docx_path = tmp_path / "fixture.docx"
        document = Document()
        document.add_heading("图片查重样本", level=1)
        document.add_paragraph("下图必须被实体化。")
        document.add_picture(str(image_path))
        document.save(docx_path)

        images_dir = tmp_path / "fixture_images"
        result = MarkitdownConverter().convert(docx_path, images_dir=images_dir)

        assert len(result.images) == 1
        assert (images_dir / result.images[0].filename).is_file()
        assert images_dir.name in result.markdown_content
        assert "data:image" not in result.markdown_content
        assert "base64" not in result.markdown_content

    def test_docx_image_inside_table_cell_keeps_markdown_ref(self, tmp_path):
        """Regression (prod 2026-08-18 永定区侨育中学项目): bid documents place
        certificate scans / verification screenshots inside (borderless)
        tables.  markdownify's default keep_inline_images_in=[] drops cell
        images keeping only alt text, so descr-less screenshots vanished from
        the parsed markdown and the review agent reported them missing.  The
        converter must keep `![alt](path)` refs inside table cells."""
        from docx import Document
        from PIL import Image

        image_path = tmp_path / "fixture.png"
        Image.new("RGB", (640, 360), "white").save(image_path)
        docx_path = tmp_path / "fixture.docx"
        document = Document()
        document.add_heading("认证体系", level=1)
        table = document.add_table(rows=1, cols=1)
        # add_picture emits an inline drawing without docPr descr → empty alt,
        # the exact shape that used to collapse to an empty table cell.
        run = table.cell(0, 0).paragraphs[0].add_run()
        run.add_picture(str(image_path))
        document.save(docx_path)

        images_dir = tmp_path / "fixture_images"
        result = MarkitdownConverter().convert(docx_path, images_dir=images_dir)

        assert len(result.images) == 1
        ref_rows = [
            line
            for line in result.markdown_content.splitlines()
            if line.startswith("|") and f"]({images_dir.name}/" in line
        ]
        assert ref_rows, (
            "image ref inside a table cell was dropped from markdown:\n"
            + result.markdown_content
        )

    def test_docx_with_null_image_relationship_converts_with_placeholder(self, tmp_path):
        """Regression (prod 2026-08-15): DOCX files from some bid-authoring
        tools carry an image relationship with Target="../NULL" whose part is
        absent from the archive.  mammoth raises KeyError("There is no item
        named 'word/../NULL' in the archive"), which used to fail the whole
        conversion; the broken image must become a 1x1 placeholder instead."""
        import re
        import zipfile
        from docx import Document
        from PIL import Image

        image_path = tmp_path / "fixture.png"
        Image.new("RGB", (320, 240), "white").save(image_path)
        docx_path = tmp_path / "fixture.docx"
        document = Document()
        document.add_heading("九州投标文件", level=1)
        document.add_paragraph("正文必须完整保留。")
        document.add_picture(str(image_path))
        document.save(docx_path)

        # Clone the existing drawing, retarget its blip at a relationship whose
        # Target="../NULL" has no part inside the archive.
        with zipfile.ZipFile(docx_path) as zin:
            entries = {name: zin.read(name) for name in zin.namelist()}
        document_xml = entries["word/document.xml"].decode("utf-8")
        drawing = re.search(r"<w:drawing>.*?</w:drawing>", document_xml, re.DOTALL).group(0)
        embed_id = re.search(r'r:embed="([^"]+)"', drawing).group(1)
        broken_drawing = drawing.replace(f'r:embed="{embed_id}"', 'r:embed="rIdNULLIMG"')
        document_xml = document_xml.replace(
            "<w:body>", "<w:body>" + f"<w:p><w:r>{broken_drawing}</w:r></w:p>", 1
        )
        entries["word/document.xml"] = document_xml.encode("utf-8")
        rels_xml = entries["word/_rels/document.xml.rels"].decode("utf-8")
        rels_xml = rels_xml.replace(
            "</Relationships>",
            '<Relationship Id="rIdNULLIMG" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" '
            'Target="../NULL"/></Relationships>',
        )
        entries["word/_rels/document.xml.rels"] = rels_xml.encode("utf-8")

        corrupted_path = tmp_path / "fixture_null.docx"
        with zipfile.ZipFile(corrupted_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for name, data in entries.items():
                zout.writestr(name, data)

        images_dir = tmp_path / "fixture_null_images"
        result = MarkitdownConverter().convert(corrupted_path, images_dir=images_dir)

        assert "正文必须完整保留" in result.markdown_content
        assert len(result.images) == 2
        sizes = {}
        for img in result.images:
            with Image.open(images_dir / img.filename) as im:
                sizes[img.filename] = im.size
            assert img.filename in result.markdown_content
        assert sorted(sizes.values()) == [(1, 1), (320, 240)]


class TestDirectFileImageHandler:
    """Test the DirectFileImageHandler that writes images directly to disk."""

    def test_handler_writes_image_to_disk(self, tmp_path):
        images_dir = tmp_path / "doc_images"
        handler = DirectFileImageHandler(images_dir, "doc_images")

        mock_image = MagicMock()
        mock_image.content_type = "image/png"
        mock_image.open.return_value.__enter__ = lambda s: MagicMock(read=lambda: b"\x89PNG\r\n\x1a\n")
        mock_image.open.return_value.__exit__ = MagicMock(return_value=False)

        result = handler(mock_image)

        assert result == {"src": "doc_images/image_1.png"}
        assert (images_dir / "image_1.png").exists()
        assert (images_dir / "image_1.png").read_bytes() == b"\x89PNG\r\n\x1a\n"
        assert len(handler.images) == 1
        assert handler.images[0].filename == "image_1.png"
        assert handler.images[0].data == b"\x89PNG\r\n\x1a\n"

    def test_handler_sequential_naming(self, tmp_path):
        images_dir = tmp_path / "doc_images"
        handler = DirectFileImageHandler(images_dir, "doc_images")

        for i in range(3):
            mock_image = MagicMock()
            mock_image.content_type = "image/jpeg"
            mock_image.open.return_value.__enter__ = lambda s: MagicMock(read=lambda: f"img{i}".encode())
            mock_image.open.return_value.__exit__ = MagicMock(return_value=False)
            handler(mock_image)

        assert len(handler.images) == 3
        assert handler.images[0].filename == "image_1.jpeg"
        assert handler.images[1].filename == "image_2.jpeg"
        assert handler.images[2].filename == "image_3.jpeg"
        assert (images_dir / "image_1.jpeg").exists()
        assert (images_dir / "image_2.jpeg").exists()
        assert (images_dir / "image_3.jpeg").exists()

    def test_handler_creates_directory(self, tmp_path):
        images_dir = tmp_path / "nested" / "images"
        handler = DirectFileImageHandler(images_dir, "nested/images")

        mock_image = MagicMock()
        mock_image.content_type = "image/png"
        mock_image.open.return_value.__enter__ = lambda s: MagicMock(read=lambda: b"data")
        mock_image.open.return_value.__exit__ = MagicMock(return_value=False)

        handler(mock_image)

        assert images_dir.exists()
        assert (images_dir / "image_1.png").exists()

    def test_handler_default_ext_for_unknown_content_type(self, tmp_path):
        images_dir = tmp_path / "doc_images"
        handler = DirectFileImageHandler(images_dir, "doc_images")

        mock_image = MagicMock()
        mock_image.content_type = None
        mock_image.open.return_value.__enter__ = lambda s: MagicMock(read=lambda: b"data")
        mock_image.open.return_value.__exit__ = MagicMock(return_value=False)

        result = handler(mock_image)

        assert result == {"src": "doc_images/image_1.png"}
        assert handler.images[0].filename == "image_1.png"

    def test_handler_returns_file_path_reference(self, tmp_path):
        """Verify handler returns a lightweight file path, NOT a data URI."""
        images_dir = tmp_path / "my_doc_images"
        handler = DirectFileImageHandler(images_dir, "my_doc_images")

        mock_image = MagicMock()
        mock_image.content_type = "image/png"
        mock_image.open.return_value.__enter__ = lambda s: MagicMock(read=lambda: b"x" * 1000)
        mock_image.open.return_value.__exit__ = MagicMock(return_value=False)

        result = handler(mock_image)

        assert "data:" not in result["src"]
        assert "base64" not in result["src"]
        assert result["src"] == "my_doc_images/image_1.png"

    def test_handler_preserves_jpx_for_later_ocr_normalization(self, tmp_path):
        images_dir = tmp_path / "doc_images"
        handler = DirectFileImageHandler(images_dir, "doc_images")
        mock_image = MagicMock()
        mock_image.content_type = "image/jpx"
        mock_image.open.return_value.__enter__ = lambda s: MagicMock(
            read=lambda: b"jpeg-2000-source"
        )
        mock_image.open.return_value.__exit__ = MagicMock(return_value=False)

        result = handler(mock_image)

        assert result == {"src": "doc_images/image_1.jpx"}
        assert (images_dir / "image_1.jpx").read_bytes() == b"jpeg-2000-source"

    def test_handler_substitutes_placeholder_when_image_part_missing(self, tmp_path):
        """word/../NULL-style broken relationships must not kill the handler."""
        images_dir = tmp_path / "doc_images"
        handler = DirectFileImageHandler(images_dir, "doc_images")

        mock_image = MagicMock()
        mock_image.content_type = "image/png"
        mock_image.open.side_effect = KeyError(
            "There is no item named 'word/../NULL' in the archive"
        )

        result = handler(mock_image)

        assert result == {"src": "doc_images/image_1.png"}
        assert handler.failed_images == 1
        assert len(handler.images) == 1
        placeholder_file = images_dir / "image_1.png"
        assert placeholder_file.is_file()
        assert placeholder_file.read_bytes() == handler.images[0].data

    def test_handler_placeholder_forces_png_ext_for_odd_content_type(self, tmp_path):
        """A broken image with a non-image content type still lands on .png."""
        images_dir = tmp_path / "doc_images"
        handler = DirectFileImageHandler(images_dir, "doc_images")

        mock_image = MagicMock()
        mock_image.content_type = "application/octet-stream"
        mock_image.open.side_effect = KeyError("There is no item named 'NULL' in the archive")

        result = handler(mock_image)

        assert result == {"src": "doc_images/image_1.png"}
        assert (images_dir / "image_1.png").is_file()

    def test_handler_failed_image_does_not_break_numbering(self, tmp_path):
        """A failed image consumes its slot; subsequent images keep numbering."""
        images_dir = tmp_path / "doc_images"
        handler = DirectFileImageHandler(images_dir, "doc_images")

        broken = MagicMock()
        broken.content_type = "image/png"
        broken.open.side_effect = KeyError("word/../NULL")
        handler(broken)

        good = MagicMock()
        good.content_type = "image/png"
        good.open.return_value.__enter__ = lambda s: MagicMock(read=lambda: b"real-data")
        good.open.return_value.__exit__ = MagicMock(return_value=False)
        result = handler(good)

        assert result == {"src": "doc_images/image_2.png"}
        assert (images_dir / "image_2.png").read_bytes() == b"real-data"
        assert [img.filename for img in handler.images] == ["image_1.png", "image_2.png"]
        assert handler.failed_images == 1
